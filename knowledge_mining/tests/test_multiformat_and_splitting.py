"""Tests for multi-format parsing (HTML, DOCX), PDF Chinese heading detection,
and segment token upper-limit splitting.

Covers:
- Phase 1: HTML → markdown conversion
- Phase 2: DOCX structural parsing
- Phase 3: PDF Chinese heading recognition + font-size heuristics
- Phase 4: Segment token limit splitting
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest

from knowledge_mining.mining.contracts.models import (
    ContentBlock,
    DocumentProfile,
    RawSegmentData,
    SectionNode,
)


# ===================================================================
# Helpers
# ===================================================================

def _collect_blocks(node: SectionNode) -> list[ContentBlock]:
    """Recursively collect all ContentBlocks from a SectionNode tree."""
    blocks = list(node.blocks)
    for child in node.children:
        blocks.extend(_collect_blocks(child))
    return blocks


def _collect_headings(node: SectionNode) -> list[ContentBlock]:
    """Collect all heading blocks from a SectionNode tree."""
    return [b for b in _collect_blocks(node) if b.block_type == "heading"]


def _collect_section_titles(node: SectionNode) -> list[str]:
    """Recursively collect all non-root section titles."""
    titles: list[str] = []
    for child in node.children:
        if child.title:
            titles.append(child.title)
        titles.extend(_collect_section_titles(child))
    return titles


def _collect_section_info(node: SectionNode) -> list[tuple[str, int]]:
    """Recursively collect (title, level) pairs from all sections."""
    info: list[tuple[str, int]] = []
    for child in node.children:
        if child.title:
            info.append((child.title, child.level))
        info.extend(_collect_section_info(child))
    return info


def _make_profile(doc_key: str = "test-doc") -> DocumentProfile:
    return DocumentProfile(document_key=doc_key)


# ===================================================================
# Phase 1: HTML → Markdown
# ===================================================================

class TestHtmlToMarkdown:
    """Tests for html_to_markdown() in preprocessing.py."""

    def test_simple_html_produces_markdown(self, tmp_path: Path):
        from knowledge_mining.mining.ingestion.preprocessing import html_to_markdown

        html_file = tmp_path / "test.html"
        html_file.write_text("""<html><body>
<h1>Main Title</h1>
<p>This is a paragraph.</p>
<h2>Section A</h2>
<p>Content of section A.</p>
</body></html>""", encoding="utf-8")

        md = html_to_markdown(html_file, doc_title="TestDoc")
        assert "# TestDoc" in md
        assert "# Main Title" in md  # h1 → #, with heading_offset=0
        assert "This is a paragraph." in md
        assert "## Section A" in md  # h2 → ##
        assert "Content of section A." in md

    def test_html_with_table(self, tmp_path: Path):
        from knowledge_mining.mining.ingestion.preprocessing import html_to_markdown

        html_file = tmp_path / "table.html"
        html_file.write_text("""<html><body>
<table>
<tr><th>Name</th><th>Value</th></tr>
<tr><td>A</td><td>1</td></tr>
</table>
</body></html>""", encoding="utf-8")

        md = html_to_markdown(html_file)
        assert "Name" in md
        assert "Value" in md
        assert "---" in md  # table separator

    def test_html_with_list(self, tmp_path: Path):
        from knowledge_mining.mining.ingestion.preprocessing import html_to_markdown

        html_file = tmp_path / "list.html"
        html_file.write_text("""<html><body>
<ul>
<li>Item A</li>
<li>Item B</li>
</ul>
</body></html>""", encoding="utf-8")

        md = html_to_markdown(html_file)
        assert "- Item A" in md
        assert "- Item B" in md

    def test_html_ingestion_integration(self, tmp_path: Path):
        from knowledge_mining.mining.ingestion import ingest_directory

        html_file = tmp_path / "doc.html"
        html_file.write_text("""<html><body>
<h1>Title</h1>
<p>Some content.</p>
</body></html>""", encoding="utf-8")

        docs, summary = ingest_directory(tmp_path)
        assert summary["parsed_documents"] == 1
        assert len(docs) == 1
        assert docs[0].file_type == "markdown"  # converted to markdown
        assert "Title" in docs[0].content
        assert docs[0].metadata_json.get("source_format") == "html"


# ===================================================================
# Phase 2: DOCX Parsing
# ===================================================================

class TestDocxParser:
    """Tests for DocxParser and parse_docx_to_section_tree."""

    @pytest.fixture()
    def sample_docx(self, tmp_path: Path) -> Path:
        """Create a minimal DOCX with headings, paragraphs, and a table."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading("Chapter One", level=1)
        doc.add_paragraph("First paragraph in chapter one.")
        doc.add_heading("Section 1.1", level=2)
        doc.add_paragraph("Content under section 1.1.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "SMF"
        table.cell(1, 1).text = "100"

        doc.add_heading("Chapter Two", level=1)
        doc.add_paragraph("Second chapter content.")

        path = tmp_path / "test.docx"
        doc.save(str(path))
        return path

    def test_docx_headings_detected(self, sample_docx: Path):
        from knowledge_mining.mining.infra.docx_parser import parse_docx_to_section_tree

        tree = parse_docx_to_section_tree(str(sample_docx), doc_title="TestDoc")
        titles = _collect_section_titles(tree)

        assert "Chapter One" in titles
        assert "Section 1.1" in titles
        assert "Chapter Two" in titles

    def test_docx_heading_levels(self, sample_docx: Path):
        from knowledge_mining.mining.infra.docx_parser import parse_docx_to_section_tree

        tree = parse_docx_to_section_tree(str(sample_docx), doc_title="TestDoc")
        info = _collect_section_info(tree)

        level_map = dict(info)
        assert level_map.get("Chapter One") == 1
        assert level_map.get("Section 1.1") == 2
        assert level_map.get("Chapter Two") == 1

    def test_docx_table_extracted(self, sample_docx: Path):
        from knowledge_mining.mining.infra.docx_parser import parse_docx_to_section_tree

        tree = parse_docx_to_section_tree(str(sample_docx), doc_title="TestDoc")
        blocks = _collect_blocks(tree)
        table_blocks = [b for b in blocks if b.block_type == "table"]

        assert len(table_blocks) >= 1
        table_text = table_blocks[0].text
        assert "Name" in table_text
        assert "SMF" in table_text

    def test_docx_section_tree_structure(self, sample_docx: Path):
        from knowledge_mining.mining.infra.docx_parser import parse_docx_to_section_tree

        tree = parse_docx_to_section_tree(str(sample_docx), doc_title="TestDoc")

        # Root should have children for Chapter One and Chapter Two
        child_titles = [c.title for c in tree.children]
        assert "Chapter One" in child_titles
        assert "Chapter Two" in child_titles

        # Chapter One should have a child for Section 1.1
        ch1 = next(c for c in tree.children if c.title == "Chapter One")
        sub_titles = [c.title for c in ch1.children]
        assert "Section 1.1" in sub_titles

    def test_docx_parser_via_parse_stage(self, sample_docx: Path):
        from knowledge_mining.mining.stages.parse import create_parser

        parser = create_parser("docx")
        tree = parser.parse("", "test.docx", {"file_path": str(sample_docx)})
        assert tree is not None
        titles = _collect_section_titles(tree)
        assert len(titles) >= 3

    def test_docx_empty_file(self, tmp_path: Path):
        from docx import Document
        from knowledge_mining.mining.infra.docx_parser import parse_docx_to_section_tree

        doc = Document()
        path = tmp_path / "empty.docx"
        doc.save(str(path))

        tree = parse_docx_to_section_tree(str(path), doc_title="Empty")
        assert tree is not None
        blocks = _collect_blocks(tree)
        assert len(blocks) == 0

    def test_docx_ingestion_integration(self, sample_docx: Path):
        from knowledge_mining.mining.ingestion import ingest_directory

        # Move docx into a tmp dir for ingestion
        ingest_dir = sample_docx.parent
        docs, summary = ingest_directory(ingest_dir)
        assert summary["parsed_documents"] == 1
        assert docs[0].metadata_json.get("source_format") in ("doc", "docx")


# ===================================================================
# Phase 3: PDF Chinese Heading Detection
# ===================================================================

class TestPdfChineseHeadings:
    """Tests for Chinese heading pattern detection in pdf_parser."""

    def test_cn_chapter_regex(self):
        from knowledge_mining.mining.infra.pdf_parser import CN_CHAPTER_RE

        assert CN_CHAPTER_RE.match("第一章 概述")
        assert CN_CHAPTER_RE.match("第二章 系统架构")
        assert CN_CHAPTER_RE.match("第三章 基本配置")
        assert CN_CHAPTER_RE.match("第12章 高级功能")
        assert not CN_CHAPTER_RE.match("这是普通段落")

    def test_cn_section_regex(self):
        from knowledge_mining.mining.infra.pdf_parser import CN_SECTION_RE

        assert CN_SECTION_RE.match("第一节 概述")
        assert CN_SECTION_RE.match("第二节 配置说明")
        assert CN_SECTION_RE.match("第3节 操作步骤")
        assert not CN_SECTION_RE.match("普通文本内容")

    def test_cn_enum_regex(self):
        from knowledge_mining.mining.infra.pdf_parser import CN_ENUM_RE

        assert CN_ENUM_RE.match("（一）总体要求")
        assert CN_ENUM_RE.match("(二)详细配置")
        assert not CN_ENUM_RE.match("普通文本")

    def test_cn_dash_enum_regex(self):
        from knowledge_mining.mining.infra.pdf_parser import CN_DASH_ENUM_RE

        assert CN_DASH_ENUM_RE.match("一、总体要求")
        assert CN_DASH_ENUM_RE.match("二、配置说明")
        assert not CN_DASH_ENUM_RE.match("普通文本")

    def test_try_cn_heading_levels(self):
        from knowledge_mining.mining.infra.pdf_parser import _try_cn_heading

        assert _try_cn_heading("第一章 概述") == 1
        assert _try_cn_heading("第二章 系统架构") == 1
        assert _try_cn_heading("第一节 概述") == 2
        assert _try_cn_heading("第二节 配置说明") == 2
        assert _try_cn_heading("（一）总体要求") == 3
        assert _try_cn_heading("一、总体要求") == 3
        assert _try_cn_heading("这是普通文本") is None

    def test_font_size_to_level(self):
        from knowledge_mining.mining.infra.pdf_parser import _font_size_to_level

        sizes = [18.0, 15.0, 13.0, 10.0]  # descending distinct sizes
        assert _font_size_to_level(18.0, sizes) == 1
        assert _font_size_to_level(15.0, sizes) == 2
        assert _font_size_to_level(13.0, sizes) == 3
        assert _font_size_to_level(10.0, sizes) == 4

    def test_cn_heading_in_classify_blocks(self):
        from knowledge_mining.mining.infra.pdf_parser import _classify_blocks, _PdfBlock

        blocks = [
            _PdfBlock(page_no=1, text="第一章 概述", font_size=16.0, x0=72.0, y0=700.0, page_height=842.0),
            _PdfBlock(page_no=1, text="这是概述段落的内容。", font_size=10.0, x0=72.0, y0=650.0, page_height=842.0),
            _PdfBlock(page_no=1, text="第二节 详细说明", font_size=14.0, x0=72.0, y0=600.0, page_height=842.0),
            _PdfBlock(page_no=1, text="详细说明的内容。", font_size=10.0, x0=72.0, y0=550.0, page_height=842.0),
        ]

        content_blocks = _classify_blocks(blocks)
        headings = [b for b in content_blocks if b.block_type == "heading"]

        heading_texts = [h.text for h in headings]
        assert "第一章 概述" in heading_texts
        assert "第二节 详细说明" in heading_texts

        # Verify heading levels
        level_map = {h.text: h.level for h in headings}
        assert level_map["第一章 概述"] == 1
        assert level_map["第二节 详细说明"] == 2

    def test_font_size_heuristic_heading(self):
        """Large font + short text should be detected as heading."""
        from knowledge_mining.mining.infra.pdf_parser import _classify_blocks, _PdfBlock

        blocks = [
            # body size = 10.0 (most common)
            _PdfBlock(page_no=1, text="大标题文本", font_size=18.0, x0=72.0, y0=700.0, page_height=842.0),
            _PdfBlock(page_no=1, text="正文段落一，这是普通内容。" * 5, font_size=10.0, x0=72.0, y0=650.0, page_height=842.0),
            _PdfBlock(page_no=1, text="正文段落二，这也是普通内容。" * 5, font_size=10.0, x0=72.0, y0=600.0, page_height=842.0),
            _PdfBlock(page_no=1, text="正文段落三，仍然是普通内容。" * 5, font_size=10.0, x0=72.0, y0=550.0, page_height=842.0),
        ]

        content_blocks = _classify_blocks(blocks)
        headings = [b for b in content_blocks if b.block_type == "heading"]

        assert len(headings) >= 1
        assert headings[0].text == "大标题文本"

    def test_split_long_blocks(self):
        from knowledge_mining.mining.infra.pdf_parser import _split_long_blocks, _PdfBlock

        # A block with multiple paragraphs separated by double newlines
        long_text = "第一段内容。\n\n第二段内容。\n\n第三段内容。"
        blocks = [
            _PdfBlock(page_no=1, text=long_text, font_size=10.0, x0=72.0, y0=700.0, page_height=842.0),
        ]

        split = _split_long_blocks(blocks)
        assert len(split) == 3
        assert split[0].text == "第一段内容。"
        assert split[1].text == "第二段内容。"
        assert split[2].text == "第三段内容。"


# ===================================================================
# Phase 4: Segment Token Limit Splitting
# ===================================================================

class TestSegmentTokenSplitting:
    """Tests for _split_large_segments in segment.py."""

    def _make_segment(self, text: str, **kwargs) -> RawSegmentData:
        from knowledge_mining.mining.infra.hash_utils import content_hash, normalized_hash
        from knowledge_mining.mining.infra.text_utils import token_count

        return RawSegmentData(
            document_key="test-doc",
            segment_index=0,
            block_type=kwargs.get("block_type", "paragraph"),
            raw_text=text,
            normalized_text=text.lower().strip(),
            content_hash=content_hash(text),
            normalized_hash=normalized_hash(text),
            token_count=token_count(text),
            section_path=kwargs.get("section_path", []),
            section_title=kwargs.get("section_title", "TestSection"),
        )

    def test_small_segments_untouched(self):
        from knowledge_mining.mining.stages.segment import _split_large_segments
        from knowledge_mining.mining.infra.text_utils import token_count

        text = "这是一段短文本。"
        seg = self._make_segment(text)
        assert token_count(text) < 512

        result = _split_large_segments([seg], max_tokens=512)
        assert len(result) == 1
        assert result[0].raw_text == text

    def test_large_segment_split_at_paragraphs(self):
        from knowledge_mining.mining.stages.segment import _split_large_segments
        from knowledge_mining.mining.infra.text_utils import token_count

        # Create text with multiple paragraphs that exceeds 512 tokens
        para = "这是一段比较长的中文文本，用于测试段落拆分功能。" * 20
        text = f"{para}\n\n{para}\n\n{para}"
        tc = token_count(text)
        assert tc > 512, f"Expected >512 tokens, got {tc}"

        seg = self._make_segment(text)
        result = _split_large_segments([seg], max_tokens=512)

        assert len(result) >= 2
        # Each sub-segment should be within token limit
        for r in result:
            assert r.token_count <= 512 + 50, f"Sub-segment has {r.token_count} tokens"  # small tolerance for boundary

    def test_split_preserves_metadata(self):
        from knowledge_mining.mining.stages.segment import _split_large_segments

        para = "这是测试文本内容。" * 20
        text = f"{para}\n\n{para}\n\n{para}\n\n{para}"
        seg = self._make_segment(
            text,
            section_path=[{"title": "Ch1", "level": 1}],
            section_title="Ch1",
        )

        result = _split_large_segments([seg], max_tokens=512)
        assert len(result) >= 2

        for r in result:
            assert r.document_key == "test-doc"
            assert r.section_path == [{"title": "Ch1", "level": 1}]
            assert r.section_title == "Ch1"

    def test_split_at_sentence_boundaries(self):
        """When a single paragraph exceeds the limit, it should be split by sentences."""
        from knowledge_mining.mining.stages.segment import _split_large_segments
        from knowledge_mining.mining.infra.text_utils import token_count

        # Single paragraph (no \n\n) but very long
        text = "这是第一个句子。这是第二个句子。这是第三个句子。这是第四个句子。这是第五个句子。" * 30
        tc = token_count(text)
        assert tc > 200

        seg = self._make_segment(text)
        result = _split_large_segments([seg], max_tokens=200)

        assert len(result) >= 2
        for r in result:
            # Each sub-segment should be roughly within limit (allow tolerance for sentence granularity)
            assert r.token_count <= 400, f"Sub-segment too large: {r.token_count}"

    def test_split_reindexes_hashes(self):
        from knowledge_mining.mining.stages.segment import _split_large_segments
        from knowledge_mining.mining.infra.hash_utils import content_hash

        para = "测试内容。" * 30
        text = f"{para}\n\n{para}\n\n{para}\n\n{para}"
        seg = self._make_segment(text)

        result = _split_large_segments([seg], max_tokens=200)
        assert len(result) >= 2

        # Each sub-segment should have correct hash
        for r in result:
            assert r.content_hash == content_hash(r.raw_text)
            assert r.token_count is not None and r.token_count > 0

    def test_mixed_small_and_large_segments(self):
        from knowledge_mining.mining.stages.segment import _split_large_segments

        small_text = "短文本。"
        large_para = "这是一段长文本内容。" * 40
        large_text = f"{large_para}\n\n{large_para}\n\n{large_para}"

        segs = [
            self._make_segment(small_text),
            self._make_segment(large_text),
            self._make_segment(small_text),
        ]

        result = _split_large_segments(segs, max_tokens=512)
        # Small segments should be preserved, large one split
        assert len(result) >= 3
        assert result[0].raw_text == small_text
        assert result[-1].raw_text == small_text

    def test_empty_segments_list(self):
        from knowledge_mining.mining.stages.segment import _split_large_segments

        result = _split_large_segments([], max_tokens=512)
        assert result == []


# ===================================================================
# Text Utils: split_sentences
# ===================================================================

class TestSplitSentences:
    """Tests for split_sentences in text_utils.py."""

    def test_cjk_sentence_split(self):
        from knowledge_mining.mining.infra.text_utils import split_sentences

        text = "第一句话。第二句话！第三句话？"
        sents = split_sentences(text)
        assert len(sents) == 3
        assert sents[0] == "第一句话。"
        assert sents[1] == "第二句话！"
        assert sents[2] == "第三句话？"

    def test_newline_split(self):
        from knowledge_mining.mining.infra.text_utils import split_sentences

        text = "第一行\n第二行\n第三行"
        sents = split_sentences(text)
        assert len(sents) == 3

    def test_mixed_punctuation(self):
        from knowledge_mining.mining.infra.text_utils import split_sentences

        text = "配置步骤如下。首先安装软件？然后进行设置！最后验证。"
        sents = split_sentences(text)
        assert len(sents) == 4

    def test_empty_text(self):
        from knowledge_mining.mining.infra.text_utils import split_sentences

        assert split_sentences("") == []
        assert split_sentences("   ") == []

    def test_no_boundary(self):
        from knowledge_mining.mining.infra.text_utils import split_sentences

        text = "这是一段没有句号的文本"
        sents = split_sentences(text)
        assert len(sents) == 1
        assert sents[0] == text
