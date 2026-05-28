"""Structural DOCX parser using python-docx.

Strategy:
1. Open the DOCX file via python-docx Document.
2. Walk paragraphs in order, detecting headings by style name (Heading 1..6).
3. Detect tables as ContentBlock(block_type="table").
4. Detect list items via paragraph style (List Paragraph) or bullet prefixes.
5. Build a SectionNode tree using the same stack-based builder as pdf_parser.
"""
from __future__ import annotations

import logging
import re
from typing import Any

from knowledge_mining.mining.contracts.models import ContentBlock, SectionNode

logger = logging.getLogger(__name__)

_HEADING_STYLE_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
_LIST_STYLE_RE = re.compile(r"^List", re.IGNORECASE)
_BULLET_PREFIX_RE = re.compile(r"^[•\-\*\u2022\u25CF\u25CB]\s+")


def parse_docx_to_section_tree(
    file_path: str, doc_title: str | None = None,
) -> SectionNode:
    """Parse a DOCX file into a SectionNode tree."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed; cannot parse DOCX files")
        return SectionNode(title=doc_title, level=0)

    doc = Document(file_path)
    blocks = _extract_blocks(doc)
    if not blocks:
        return SectionNode(title=doc_title, level=0)
    return _build_section_tree(blocks, doc_title)


def _extract_blocks(doc: Any) -> list[ContentBlock]:
    """Walk document body elements in order, yielding ContentBlocks."""
    blocks: list[ContentBlock] = []

    # doc.element.body contains both paragraphs and tables in document order.
    from docx.oxml.ns import qn

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            paragraph = None
            for p in doc.paragraphs:
                if p._element is element:
                    paragraph = p
                    break
            if paragraph is None:
                continue
            block = _classify_paragraph(paragraph)
            if block is not None:
                blocks.append(block)

        elif tag == "tbl":
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table is not None:
                blocks.append(_extract_table(table))

    return blocks


def _classify_paragraph(para: Any) -> ContentBlock | None:
    """Classify a paragraph as heading, list item, or plain paragraph."""
    text = para.text.strip()
    if not text:
        return None

    style_name = (para.style.name or "") if para.style else ""

    # Check heading styles
    m = _HEADING_STYLE_RE.match(style_name)
    if m:
        level = int(m.group(1))
        level = max(1, min(6, level))
        return ContentBlock(block_type="heading", text=text, level=level)

    # Check list styles
    is_list = _LIST_STYLE_RE.match(style_name) is not None if style_name else False
    if is_list or _BULLET_PREFIX_RE.match(text):
        return ContentBlock(block_type="list", text=text)

    return ContentBlock(block_type="paragraph", text=text)


def _extract_table(table: Any) -> ContentBlock:
    """Extract a table into a ContentBlock with pipe-delimited text."""
    rows_text: list[str] = []
    col_count = 0
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        col_count = max(col_count, len(cells))
        rows_text.append("| " + " | ".join(cells) + " |")

    if not rows_text:
        return ContentBlock(block_type="table", text="")

    # Add separator after first row (header)
    if len(rows_text) >= 2:
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows_text.insert(1, separator)

    table_text = "\n".join(rows_text)
    return ContentBlock(
        block_type="table",
        text=table_text,
        structure={"kind": "table", "row_count": len(table.rows), "col_count": col_count},
    )


def _build_section_tree(
    blocks: list[ContentBlock], doc_title: str | None,
) -> SectionNode:
    """Stack-based builder: pop ancestors with level >= current heading level."""
    root = _mutable_node(title=doc_title, level=0)
    stack: list[dict] = [root]

    for block in blocks:
        if block.block_type == "heading" and block.level:
            level = block.level
            while len(stack) > 1 and stack[-1]["level"] >= level:
                stack.pop()
            new_node = _mutable_node(title=block.text, level=level)
            stack[-1]["children"].append(new_node)
            stack.append(new_node)
        else:
            stack[-1]["blocks"].append(block)

    return _freeze(root)


def _mutable_node(title: str | None, level: int) -> dict:
    return {"title": title, "level": level, "children": [], "blocks": []}


def _freeze(node: dict) -> SectionNode:
    return SectionNode(
        title=node["title"],
        level=node["level"],
        blocks=tuple(node["blocks"]),
        children=tuple(_freeze(c) for c in node["children"]),
    )
